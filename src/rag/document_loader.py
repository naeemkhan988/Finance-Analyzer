"""
Document Loader & Processor
============================
Handles loading, parsing, and chunking of financial documents.
Supports: PDF, DOCX, TXT, XLSX, CSV
Includes: OCR support, table extraction, metadata enrichment.
"""

import os
import re
import csv
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from datetime import datetime

import pdfplumber
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class DocumentChunk:
    """Represents a chunk of a document with metadata."""

    def __init__(
        self,
        content: str,
        metadata: Dict,
        chunk_id: str = "",
        embedding: Optional[List[float]] = None,
    ):
        self.content = content
        self.metadata = metadata
        self.chunk_id = chunk_id or self._generate_id()
        self.embedding = embedding

    def _generate_id(self) -> str:
        """Generate unique chunk ID from content hash."""
        return hashlib.md5(self.content.encode()).hexdigest()[:12]

    def to_dict(self) -> Dict:
        return {
            "content": self.content,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id,
        }


class DocumentProcessor:
    """
    Production-grade document processor with multi-format support.

    Features:
    - Smart chunking with configurable overlap
    - Table extraction from PDFs
    - Metadata enrichment (page numbers, source, timestamps)
    - Content deduplication
    - Financial entity detection
    """

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 128):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._financial_patterns = self._compile_financial_patterns()

    def _compile_financial_patterns(self) -> Dict[str, re.Pattern]:
        """Compile regex patterns for financial entity detection."""
        return {
            "currency": re.compile(
                r"\$[\d,]+\.?\d*|\€[\d,]+\.?\d*|£[\d,]+\.?\d*"
            ),
            "percentage": re.compile(r"\d+\.?\d*\s*%"),
            "date": re.compile(
                r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
                r"[a-z]*\.?\s+\d{1,2},?\s+\d{4}\b|\b\d{1,2}/\d{1,2}/\d{2,4}\b"
            ),
            "ticker": re.compile(r"\b[A-Z]{1,5}\b(?=\s*[\(\$])"),
            "fiscal_term": re.compile(
                r"\b(?:Q[1-4]|FY\d{2,4}|revenue|earnings|EBITDA|net income|"
                r"gross margin|operating income|EPS|P/E|ROI|ROE|ROA)\b",
                re.IGNORECASE,
            ),
        }

    def load_pdf(self, file_path: str) -> List[Dict]:
        """
        Load PDF with text and table extraction.

        Returns list of page-level content with metadata.
        """
        file_path = str(file_path)
        pages = []
        file_name = Path(file_path).name
        file_hash = self._compute_file_hash(file_path)

        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    tables = self._extract_tables(page)

                    # Combine text and table content
                    content = page_text
                    if tables:
                        content += "\n\n[TABLES]\n" + "\n".join(tables)

                    if content.strip():
                        pages.append(
                            {
                                "content": content,
                                "metadata": {
                                    "source": file_name,
                                    "file_hash": file_hash,
                                    "page": i + 1,
                                    "total_pages": total_pages,
                                    "type": "pdf",
                                    "has_tables": len(tables) > 0,
                                    "table_count": len(tables),
                                    "char_count": len(content),
                                    "loaded_at": datetime.utcnow().isoformat(),
                                },
                            }
                        )

            logger.info(
                f"Loaded PDF '{file_name}': {len(pages)} pages, "
                f"{total_pages} total pages"
            )

        except Exception as e:
            logger.error(f"Error loading PDF '{file_path}': {e}")
            # Fallback to PyPDF2
            pages = self._load_pdf_fallback(file_path)

        return pages

    def _load_pdf_fallback(self, file_path: str) -> List[Dict]:
        """Fallback PDF loader using PyPDF2."""
        pages = []
        file_name = Path(file_path).name
        file_hash = self._compute_file_hash(file_path)

        try:
            reader = PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(
                        {
                            "content": text,
                            "metadata": {
                                "source": file_name,
                                "file_hash": file_hash,
                                "page": i + 1,
                                "total_pages": len(reader.pages),
                                "type": "pdf",
                                "has_tables": False,
                                "loaded_at": datetime.utcnow().isoformat(),
                            },
                        }
                    )
            logger.info(f"Fallback loaded PDF '{file_name}': {len(pages)} pages")
        except Exception as e:
            logger.error(f"Fallback PDF loading failed for '{file_path}': {e}")

        return pages

    def _extract_tables(self, page) -> List[str]:
        """Extract tables from a pdfplumber page and format as text."""
        tables = []
        try:
            extracted = page.extract_tables()
            if extracted:
                for table in extracted:
                    if table:
                        formatted = self._format_table(table)
                        if formatted:
                            tables.append(formatted)
        except Exception as e:
            logger.warning(f"Table extraction failed: {e}")
        return tables

    def _format_table(self, table: List[List]) -> str:
        """Format extracted table as readable text."""
        if not table or not table[0]:
            return ""

        rows = []
        headers = [str(h or "").strip() for h in table[0]]
        rows.append(" | ".join(headers))
        rows.append("-" * len(rows[0]))

        for row in table[1:]:
            cells = [str(c or "").strip() for c in row]
            rows.append(" | ".join(cells))

        return "\n".join(rows)

    def load_text(self, file_path: str) -> List[Dict]:
        """Load plain text file."""
        file_name = Path(file_path).name
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return [
                {
                    "content": content,
                    "metadata": {
                        "source": file_name,
                        "file_hash": self._compute_file_hash(file_path),
                        "page": 1,
                        "total_pages": 1,
                        "type": "txt",
                        "loaded_at": datetime.utcnow().isoformat(),
                    },
                }
            ]
        except Exception as e:
            logger.error(f"Error loading text file '{file_path}': {e}")
            return []

    def load_docx(self, file_path: str) -> List[Dict]:
        """Load DOCX file using python-docx."""
        file_name = Path(file_path).name
        file_hash = self._compute_file_hash(file_path)
        pages = []

        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)

            # Extract all paragraph text
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)

            # Extract tables
            table_texts = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    table_texts.append("\n".join(rows))

            if table_texts:
                full_text += "\n\n[TABLES]\n" + "\n\n".join(table_texts)

            if full_text.strip():
                # Split into logical pages (~3000 chars each for long docs)
                page_size = 3000
                if len(full_text) <= page_size:
                    page_contents = [full_text]
                else:
                    page_contents = []
                    for i in range(0, len(full_text), page_size):
                        chunk = full_text[i : i + page_size]
                        if chunk.strip():
                            page_contents.append(chunk.strip())

                for i, content in enumerate(page_contents):
                    pages.append(
                        {
                            "content": content,
                            "metadata": {
                                "source": file_name,
                                "file_hash": file_hash,
                                "page": i + 1,
                                "total_pages": len(page_contents),
                                "type": "docx",
                                "has_tables": len(table_texts) > 0,
                                "table_count": len(table_texts),
                                "char_count": len(content),
                                "loaded_at": datetime.utcnow().isoformat(),
                            },
                        }
                    )

            logger.info(
                f"Loaded DOCX '{file_name}': {len(pages)} pages, "
                f"{len(paragraphs)} paragraphs, {len(table_texts)} tables"
            )

        except ImportError:
            logger.error(
                "python-docx is not installed. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error loading DOCX '{file_path}': {e}")

        return pages

    def load_xlsx(self, file_path: str) -> List[Dict]:
        """Load XLSX file using openpyxl."""
        file_name = Path(file_path).name
        file_hash = self._compute_file_hash(file_path)
        pages = []

        try:
            from openpyxl import load_workbook

            wb = load_workbook(file_path, read_only=True, data_only=True)

            for sheet_idx, sheet_name in enumerate(wb.sheetnames):
                ws = wb[sheet_name]
                rows = []
                headers = None

                for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                    # Convert cell values to strings, handle None
                    cells = [
                        str(cell) if cell is not None else ""
                        for cell in row
                    ]
                    # Skip completely empty rows
                    if not any(c.strip() for c in cells):
                        continue

                    if headers is None:
                        headers = cells
                        rows.append(" | ".join(cells))
                        rows.append("-" * len(rows[0]))
                    else:
                        rows.append(" | ".join(cells))

                content = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)

                if content.strip() and rows:
                    pages.append(
                        {
                            "content": content,
                            "metadata": {
                                "source": file_name,
                                "file_hash": file_hash,
                                "page": sheet_idx + 1,
                                "total_pages": len(wb.sheetnames),
                                "type": "xlsx",
                                "sheet_name": sheet_name,
                                "has_tables": True,
                                "table_count": 1,
                                "row_count": len(rows) - 1,
                                "char_count": len(content),
                                "loaded_at": datetime.utcnow().isoformat(),
                            },
                        }
                    )

            wb.close()
            logger.info(
                f"Loaded XLSX '{file_name}': {len(pages)} sheets"
            )

        except ImportError:
            logger.error(
                "openpyxl is not installed. "
                "Install it with: pip install openpyxl"
            )
        except Exception as e:
            logger.error(f"Error loading XLSX '{file_path}': {e}")

        return pages

    def load_csv(self, file_path: str) -> List[Dict]:
        """Load CSV file."""
        file_name = Path(file_path).name
        file_hash = self._compute_file_hash(file_path)
        pages = []

        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                rows = []
                headers = None

                for row_idx, row in enumerate(reader):
                    cells = [cell.strip() for cell in row]
                    if not any(cells):
                        continue

                    if headers is None:
                        headers = cells
                        rows.append(" | ".join(cells))
                        rows.append("-" * len(rows[0]))
                    else:
                        rows.append(" | ".join(cells))

                content = "\n".join(rows)

                if content.strip():
                    pages.append(
                        {
                            "content": content,
                            "metadata": {
                                "source": file_name,
                                "file_hash": file_hash,
                                "page": 1,
                                "total_pages": 1,
                                "type": "csv",
                                "has_tables": True,
                                "table_count": 1,
                                "row_count": len(rows) - 1,
                                "char_count": len(content),
                                "loaded_at": datetime.utcnow().isoformat(),
                            },
                        }
                    )

            logger.info(
                f"Loaded CSV '{file_name}': {len(rows) - 1 if rows else 0} rows"
            )

        except Exception as e:
            logger.error(f"Error loading CSV '{file_path}': {e}")

        return pages

    def load_document(self, file_path: str) -> List[Dict]:
        """Load document based on file extension."""
        ext = Path(file_path).suffix.lower()

        loaders = {
            ".pdf": self.load_pdf,
            ".txt": self.load_text,
            ".docx": self.load_docx,
            ".xlsx": self.load_xlsx,
            ".csv": self.load_csv,
        }

        loader = loaders.get(ext)
        if loader:
            return loader(file_path)

        logger.warning(f"Unsupported file type: {ext}")
        return []

    def chunk_documents(self, pages: List[Dict]) -> List[DocumentChunk]:
        """
        Split documents into overlapping chunks with metadata preservation.

        Uses recursive character splitting with intelligent boundaries:
        - Paragraph breaks
        - Sentence boundaries
        - Word boundaries
        """
        chunks = []
        separators = ["\n\n", "\n", ". ", " ", ""]

        for page_data in pages:
            content = page_data["content"]
            metadata = page_data["metadata"]

            page_chunks = self._recursive_split(content, separators)

            for i, chunk_text in enumerate(page_chunks):
                if not chunk_text.strip():
                    continue

                # Detect financial entities in chunk
                financial_entities = self._detect_financial_entities(chunk_text)

                chunk_metadata = {
                    **metadata,
                    "chunk_index": i,
                    "total_chunks_in_page": len(page_chunks),
                    "financial_entities": financial_entities,
                    "word_count": len(chunk_text.split()),
                }

                chunks.append(
                    DocumentChunk(content=chunk_text.strip(), metadata=chunk_metadata)
                )

        logger.info(
            f"Created {len(chunks)} chunks from {len(pages)} pages "
            f"(chunk_size={self.chunk_size}, overlap={self.chunk_overlap})"
        )
        return chunks

    def _recursive_split(
        self, text: str, separators: List[str]
    ) -> List[str]:
        """Recursively split text using multiple separators."""
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        # Find the best separator
        sep = separators[0] if separators else ""
        chunks = []
        current_chunk = ""

        parts = text.split(sep) if sep else list(text)

        for part in parts:
            test_chunk = current_chunk + (sep if current_chunk else "") + part

            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
            else:
                if current_chunk:
                    chunks.append(current_chunk)

                    # Add overlap from previous chunk
                    if self.chunk_overlap > 0:
                        overlap_text = current_chunk[-self.chunk_overlap :]
                        current_chunk = overlap_text + sep + part
                    else:
                        current_chunk = part
                else:
                    # Part itself is too long, recurse with next separator
                    if len(separators) > 1:
                        sub_chunks = self._recursive_split(
                            part, separators[1:]
                        )
                        chunks.extend(sub_chunks)
                        current_chunk = ""
                    else:
                        # Force split at chunk_size
                        for j in range(0, len(part), self.chunk_size):
                            chunks.append(part[j : j + self.chunk_size])
                        current_chunk = ""

        if current_chunk.strip():
            chunks.append(current_chunk)

        return chunks

    def _detect_financial_entities(self, text: str) -> Dict[str, List[str]]:
        """Detect financial entities in text."""
        entities = {}
        for entity_type, pattern in self._financial_patterns.items():
            matches = pattern.findall(text)
            if matches:
                entities[entity_type] = list(set(matches[:10]))
        return entities

    def _compute_file_hash(self, file_path: str) -> str:
        """Compute SHA-256 hash of file for deduplication."""
        sha256 = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for block in iter(lambda: f.read(4096), b""):
                    sha256.update(block)
            return sha256.hexdigest()[:16]
        except Exception:
            return "unknown"

    def get_document_summary(self, pages: List[Dict]) -> Dict:
        """Generate summary statistics for a loaded document."""
        total_chars = sum(len(p["content"]) for p in pages)
        total_tables = sum(p["metadata"].get("table_count", 0) for p in pages)

        all_text = " ".join(p["content"] for p in pages)
        financial_entities = self._detect_financial_entities(all_text)

        return {
            "total_pages": len(pages),
            "total_characters": total_chars,
            "total_tables": total_tables,
            "has_financial_data": bool(financial_entities),
            "detected_entities": {
                k: len(v) for k, v in financial_entities.items()
            },
            "source": pages[0]["metadata"]["source"] if pages else "unknown",
        }
